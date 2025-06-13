import os
import re
from docx import Document
from docx.shared import Pt
from docx.opc.exceptions import PackageNotFoundError
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def remove_tags_from_text(text):
    """Remove tags like {123}, {123>, or <123} from text"""
    pattern = r'\{[0-9]{1,4}\}|\{[0-9]{1,4}\>|\<[0-9]{1,4}\}'
    return re.sub(pattern, '', text)

def set_table_borders(table):
    """Set borders for all cells in a table"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Set table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    
    # Apply borders to table
    tblPr.append(tblBorders)

def process_docx(file_path, output_path):
    """Process a single DOCX file"""
    try:
        doc = Document(file_path)
    except PackageNotFoundError:
        print(f"Skipping locked/temporary file: {os.path.basename(file_path)}")
        return
    
    # Delete first two tables if they exist
    for _ in range(2):
        if len(doc.tables) > 0:
            doc.tables[0]._element.getparent().remove(doc.tables[0]._element)
    
    # If no tables left, save and return
    if len(doc.tables) == 0:
        doc.save(output_path)
        return
    
    # STEP 1: Collect all data from all tables
    all_data = []
    for table in doc.tables:
        for row in table.rows:
            row_data = [remove_tags_from_text(cell.text) for cell in row.cells]
            all_data.append(row_data)
    
    """# Diagnostic output to verify column structure
    print(f"\nFile: {os.path.basename(file_path)}")
    print("First 3 rows with ALL columns (for verification):")
    for i, row in enumerate(all_data[:3]):
        print(f"Row {i+1}: {row}")"""
    
    # STEP 2: Create new table with correct columns
    final_table = doc.add_table(rows=0, cols=2)
    final_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row in enumerate(all_data):
        new_row = final_table.add_row()
        
        # For header row (row 1), use columns 4 and 5 (indexes 3 and 4)
        if i == 0:
            col1 = row[3] if len(row) > 3 else ""
            col2 = row[4] if len(row) > 4 else ""
        # For all other rows, use columns 4 and 6 (indexes 3 and 5)
        else:
            col1 = row[3] if len(row) > 3 else ""
            col2 = row[5] if len(row) > 5 else ""
        
        new_row.cells[0].text = col1
        new_row.cells[1].text = col2
        
        # Set vertical alignment
        for cell in new_row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    
    # Remove all original tables
    for table in list(doc.tables[:-1]):
        table._element.getparent().remove(table._element)
    
    # Format the final table
    set_table_borders(final_table)
    
    # Save the processed document
    try:
        doc.save(output_path)
        print("Processing completed successfully")
    except PermissionError:
        print(f"Could not save {output_path} - file may be open in another program")

def process_folder(folder_path):
    """Process all DOCX files in a folder"""
    for filename in os.listdir(folder_path):
        # Skip temporary Word files (start with ~$)
        if filename.startswith('~$'):
            continue
            
        if filename.endswith(".docx"):
            input_path = os.path.join(folder_path, filename)
            output_path = os.path.join(folder_path, f"{os.path.splitext(os.path.basename(filename))[0]}_processed.docx")
                
            try:
                print(f"\nProcessing: {filename}")
                process_docx(input_path, output_path)
            except Exception as e:
                print(f"Error processing {filename}: {str(e)}")

if __name__ == "__main__":
    print("Phrase Bilingual DOCX Formatting and Tag Remover")
    folder_path = input("Enter folder path containing bilingual DOCX files: ").strip('"')
    
    if os.path.isdir(folder_path):
        process_folder(folder_path)
        print("\nProcessing complete!")
    else:
        print("Error: Invalid directory path")